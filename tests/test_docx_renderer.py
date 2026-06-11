import os
import tempfile

import pytest
from docx import Document

from academic_pe.tools.docx_renderer import render_paper, set_font_style, set_paragraph_format


@pytest.fixture
def tmp_doc():
    with tempfile.TemporaryDirectory() as d:
        yield os.path.join(d, "test.docx")


class TestRenderPaper:
    def test_creates_file(self, tmp_doc):
        result = render_paper({"theory": "Hello world"}, tmp_doc)
        assert result == tmp_doc
        assert os.path.exists(tmp_doc)

    def test_all_sections_present(self, tmp_doc):
        content = {
            "theory": "Theory content here.",
            "calculation": "Calculation content here.",
            "conclusion": "Conclusion content here.",
        }
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Theory content here." in full_text
        assert "Calculation content here." in full_text
        assert "Conclusion content here." in full_text

    def test_missing_keys_rendered(self, tmp_doc):
        content = {"theory": "Only theory."}
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Only theory." in full_text

    def test_empty_content(self, tmp_doc):
        render_paper({}, tmp_doc)
        assert os.path.exists(tmp_doc)
        doc = Document(tmp_doc)
        assert len(doc.paragraphs) > 0

    def test_empty_section_string(self, tmp_doc):
        content = {"theory": "", "calculation": "  ", "conclusion": "OK"}
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "OK" in full_text

    def test_heading_rendered(self, tmp_doc):
        content = {"theory": "# My Heading\n\nParagraph text."}
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "My Heading" in full_text

    def test_bold_text(self, tmp_doc):
        content = {"theory": "This is **bold** text."}
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        found_bold = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run.bold and "bold" in run.text:
                    found_bold = True
        assert found_bold

    def test_italic_text(self, tmp_doc):
        content = {"theory": "This is *italic* text."}
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        found_italic = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run.italic and "italic" in run.text:
                    found_italic = True
        assert found_italic

    def test_latex_inline_math(self, tmp_doc):
        content = {"calculation": r"Formula $E=mc^2$ here."}
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "E=mc^2" in full_text or "E=mc2" in full_text

    def test_latex_display_math(self, tmp_doc):
        content = {"calculation": r"Display $$\sum_{i=1}^{n} i$$ end."}
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "sum" in full_text.lower() or r"\sum" in full_text

    def test_subscript_rendered(self, tmp_doc):
        content = {"calculation": r"Variable $x_{1}$ and $x_2$."}
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        found_sub = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run.font.subscript:
                    found_sub = True
        assert found_sub

    def test_times_new_roman_font(self, tmp_doc):
        content = {"theory": "Check font."}
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text.strip():
                    assert run.font.name == "Times New Roman"

    def test_justified_alignment(self, tmp_doc):
        content = {"theory": "Paragraph for alignment."}
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        found_justified = False
        for p in doc.paragraphs:
            if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                found_justified = True
        assert found_justified

    def test_default_filename(self):
        with tempfile.TemporaryDirectory() as d:
            old_cwd = os.getcwd()
            os.chdir(d)
            try:
                result = render_paper({"theory": "test"})
                assert os.path.exists(result)
            finally:
                os.chdir(old_cwd)

    def test_page_break_after_title(self, tmp_doc):
        content = {"theory": "Content."}
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        found_break = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run._element.xml.find("w:br") != -1 and 'type="page"' in run._element.xml:
                    found_break = True
        assert found_break


class TestSetFontStyle:
    def test_sets_font_name(self):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        set_font_style(run, font_name="Arial", font_size=12)
        assert run.font.name == "Arial"
        assert run.font.size is not None
        assert run.font.size.pt == 12

    def test_sets_bold(self):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        set_font_style(run, bold=True)
        assert run.bold is True

    def test_sets_italic(self):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        set_font_style(run, italic=True)
        assert run.italic is True

    def test_sets_subscript(self):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        set_font_style(run, subscript=True)
        assert run.font.subscript is True


class TestSetParagraphFormat:
    def test_sets_alignment(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_sets_line_spacing(self):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        set_paragraph_format(p, line_spacing=2.0)
        assert p.paragraph_format.line_spacing == 2.0


class TestNewRendererFeatures:
    def test_render_with_custom_style_and_sections(self, tmp_doc):
        from academic_pe.core.config import AppConfig, StyleConfig, PipelineConfig, SectionPrompt
        config = AppConfig(
            agents={},
            style=StyleConfig(
                font_name="Arial",
                font_size=12,
                title_font_size=18,
                line_spacing=2.0,
                first_line_indent_cm=1.5,
                alignment="center"
            ),
            pipeline=PipelineConfig(
                title="MY CUSTOM TITLE",
                sections=[
                    SectionPrompt(name="conclusion", topic="Done", instruction="test"),
                    SectionPrompt(name="theory", topic="Intro", instruction="test")
                ]
            )
        )
        content = {
            "theory": "Theory text.",
            "conclusion": "Conclusion text."
        }
        render_paper(content, tmp_doc, config=config)
        assert os.path.exists(tmp_doc)
        doc = Document(tmp_doc)
        
        # Verify title page
        assert doc.paragraphs[0].text.strip() == "MY CUSTOM TITLE"
        
        # Verify section order conclusion then theory
        text_runs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # First non-title text should be Conclusion text, then Theory text
        non_title_texts = [t for t in text_runs if t != "MY CUSTOM TITLE"]
        assert non_title_texts[0] == "Conclusion text."
        assert non_title_texts[1] == "Theory text."

    def test_renders_table_from_markdown(self, tmp_doc):
        content = {
            "theory": "Heading\n\n| H1 | H2 |\n|---|---|\n| V1 | V2 |\n"
        }
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 2
        assert doc.tables[0].rows[0].cells[0].text == "H1"
        assert doc.tables[0].rows[1].cells[1].text == "V2"

    def test_renders_chart_placeholder(self, tmp_doc):
        content = {
            "calculation": "Calculations\n\n[Chart: Test Performance]\n"
        }
        render_paper(content, tmp_doc)
        doc = Document(tmp_doc)
        # Check if there is an image in the document shapes or inline shapes
        assert len(doc.inline_shapes) == 1
