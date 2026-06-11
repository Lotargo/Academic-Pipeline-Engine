from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
from typing import Dict


def set_font_style(run, font_name='Times New Roman', font_size=14, bold=False, italic=False, subscript=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    run.font.subscript = subscript


def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent_cm=1.25, line_spacing=1.5, space_after=0):
    paragraph.alignment = alignment
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    paragraph_format.line_spacing = line_spacing
    paragraph_format.space_after = Pt(space_after)


def parse_math_content(paragraph, text, bold=False, italic=False):
    text = text.replace(r'\times', '×')
    pattern = re.compile(r'(_\{[^}]+\})|(_.)')
    last_pos = 0
    while True:
        match = pattern.search(text, pos=last_pos)
        if not match:
            remaining = text[last_pos:]
            if remaining:
                run = paragraph.add_run(remaining)
                set_font_style(run, bold=bold, italic=italic)
            break
        if match.start() > last_pos:
            pre = text[last_pos:match.start()]
            run = paragraph.add_run(pre)
            set_font_style(run, bold=bold, italic=italic)
        group = match.group()
        if group.startswith('_{'):
            sub_text = group[2:-1]
        else:
            sub_text = group[1:]
        run = paragraph.add_run(sub_text)
        set_font_style(run, bold=bold, italic=italic, subscript=True)
        last_pos = match.end()


def add_formatted_text(paragraph, text):
    bold_parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for b_part in bold_parts:
        if not b_part: continue
        is_bold = False
        content_to_process = b_part
        if b_part.startswith('**') and b_part.endswith('**'):
            is_bold = True
            content_to_process = b_part[2:-2]
        italic_parts = re.split(r'(\*[^*]+\*)', content_to_process)
        for i_part in italic_parts:
            if not i_part: continue
            is_italic = False
            inner_content = i_part
            if i_part.startswith('*') and i_part.endswith('*'):
                is_italic = True
                inner_content = i_part[1:-1]
            math_parts = re.split(r'(\$\$[^$]+\$\$|\$[^$]+\$)', inner_content)
            for m_part in math_parts:
                if not m_part: continue
                if m_part.startswith('$'):
                    math_content = m_part.strip('$')
                    parse_math_content(paragraph, math_content, bold=is_bold, italic=is_italic)
                else:
                    run = paragraph.add_run(m_part)
                    set_font_style(run, bold=is_bold, italic=is_italic)


def render_paper(content: Dict[str, str], output_filename: str = "Output.docx"):
    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GENERATED ACADEMIC PAPER\n")
    set_font_style(run, font_size=20, bold=True)
    doc.add_page_break()

    order = ['theory', 'calculation', 'conclusion']

    for section in order:
        if section in content:
            text_block = content[section]
            lines = text_block.split('\n')

            for line in lines:
                stripped = line.strip()
                if not stripped: continue

                if stripped.startswith('# '):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(12)
                    run = p.add_run(stripped[2:])
                    set_font_style(run, font_size=14, bold=True)
                else:
                    p = doc.add_paragraph()
                    set_paragraph_format(p)
                    add_formatted_text(p, stripped)

    doc.save(output_filename)
    return output_filename
