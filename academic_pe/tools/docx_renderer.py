from __future__ import annotations

import logging
import os
import re
import tempfile
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)


_INLINE_TOKEN_RE = re.compile(r"(\$\$.*?\$\$|\$.*?\$|\\\(.*?\\\)|\\\[.*?\\\]|\*\*.*?\*\*|\*[^*\n]+\*)", re.DOTALL)
_LIST_RE = re.compile(r"^(\s*)([-*]|\d+[.)])\s+(.+)$")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")

_LATEX_SYMBOLS = {
    r"\alpha": "alpha",
    r"\beta": "beta",
    r"\gamma": "gamma",
    r"\delta": "delta",
    r"\epsilon": "epsilon",
    r"\lambda": "lambda",
    r"\mu": "mu",
    r"\pi": "pi",
    r"\sigma": "sigma",
    r"\sum": "sum",
    r"\prod": "prod",
    r"\int": "int",
    r"\infty": "infinity",
    r"\leq": "<=",
    r"\geq": ">=",
    r"\neq": "!=",
    r"\approx": "~",
    r"\times": "x",
    r"\cdot": "*",
    r"\rightarrow": "->",
    r"\to": "->",
}


def set_font_style(
    run,
    font_name: str = "Times New Roman",
    font_size: int = 14,
    bold: bool = False,
    italic: bool = False,
    subscript: bool = False,
    superscript: bool = False,
    color: Optional[RGBColor] = None,
):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    run.font.subscript = subscript
    run.font.superscript = superscript
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(
    paragraph,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent_cm: float = 1.25,
    line_spacing: float = 1.5,
    space_after: int = 6,
):
    paragraph.alignment = alignment
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    paragraph_format.line_spacing = line_spacing
    paragraph_format.space_after = Pt(space_after)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _normalize_latex(math_content: str) -> str:
    text = math_content.strip()
    text = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", text)
    text = re.sub(r"\\sqrt\{([^{}]+)\}", r"sqrt(\1)", text)
    text = re.sub(r"\\text\{([^{}]+)\}", r"\1", text)
    for latex, replacement in _LATEX_SYMBOLS.items():
        text = text.replace(latex, replacement)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def parse_math_content(
    paragraph,
    text: str,
    font_name: str = "Times New Roman",
    font_size: int = 14,
    bold: bool = False,
    italic: bool = False,
):
    text = _normalize_latex(text)
    pattern = re.compile(r"([_^])\{?([A-Za-z0-9+\-=()]+)\}?")
    last_pos = 0
    for match in pattern.finditer(text):
        if match.start() > last_pos:
            run = paragraph.add_run(text[last_pos:match.start()])
            set_font_style(run, font_name, font_size, bold=bold, italic=italic)
        marker, value = match.groups()
        run = paragraph.add_run(value)
        set_font_style(
            run,
            font_name,
            font_size,
            bold=bold,
            italic=italic,
            subscript=marker == "_",
            superscript=marker == "^",
        )
        last_pos = match.end()
    if last_pos < len(text):
        run = paragraph.add_run(text[last_pos:])
        set_font_style(run, font_name, font_size, bold=bold, italic=italic)


def add_formatted_text(paragraph, text: str, font_name: str = "Times New Roman", font_size: int = 14):
    for part in _INLINE_TOKEN_RE.split(text):
        if not part:
            continue
        if part.startswith("$$") and part.endswith("$$"):
            parse_math_content(paragraph, part[2:-2], font_name=font_name, font_size=font_size, italic=True)
        elif part.startswith("$") and part.endswith("$"):
            parse_math_content(paragraph, part[1:-1], font_name=font_name, font_size=font_size, italic=True)
        elif part.startswith("\\(") and part.endswith("\\)"):
            parse_math_content(paragraph, part[2:-2], font_name=font_name, font_size=font_size, italic=True)
        elif part.startswith("\\[") and part.endswith("\\]"):
            parse_math_content(paragraph, part[2:-2], font_name=font_name, font_size=font_size, italic=True)
        elif part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            set_font_style(run, font_name=font_name, font_size=font_size, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            set_font_style(run, font_name=font_name, font_size=font_size, italic=True)
        else:
            run = paragraph.add_run(part)
            set_font_style(run, font_name=font_name, font_size=font_size)


def _add_markdown_paragraph(
    doc: Document,
    text: str,
    style_name: str,
    font_name: str,
    font_size: int,
    alignment,
    first_line_indent_cm: float,
    line_spacing: float,
):
    paragraph = doc.add_paragraph(style=style_name)
    if style_name == "Body Text":
        set_paragraph_format(
            paragraph,
            alignment=alignment,
            first_line_indent_cm=first_line_indent_cm,
            line_spacing=line_spacing,
        )
    add_formatted_text(paragraph, text, font_name=font_name, font_size=font_size)
    return paragraph


def _configure_styles(doc: Document, font_name: str, font_size: int, title_font_size: int) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.font.size = Pt(font_size)

    if "Body Text" not in styles:
        body = styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        body = styles["Body Text"]
    body.base_style = normal
    body.font.name = font_name
    body._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    body.font.size = Pt(font_size)
    body.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        style = styles[f"Heading {level}"]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.bold = True
        style.font.color.rgb = RGBColor(17, 94, 89)
        style.font.size = Pt(max(font_size + 5 - level, font_size))
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    if "Document Title" not in styles:
        title = styles.add_style("Document Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title = styles["Document Title"]
    title.font.name = font_name
    title._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    title.font.size = Pt(title_font_size)
    title.font.bold = True
    title.font.color.rgb = RGBColor(15, 118, 110)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)


def create_chart_image(output_path: str):
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        fig, ax = plt.subplots(figsize=(6, 3.5))
        x = [1, 2, 3, 4, 5]
        y1 = [10, 15, 25, 40, 55]
        y2 = [8, 12, 18, 28, 38]
        ax.plot(x, y1, marker="o", color="#0d9488", linewidth=2.5, label="Writer Agent")
        ax.plot(x, y2, marker="s", color="#0f766e", linewidth=2, linestyle="--", label="Reviewer Agent")
        ax.set_title("Agent Performance Metrics", fontsize=12, fontweight="bold", color="#1f2937")
        ax.set_xlabel("Iterations / Run ID", fontsize=10, color="#4b5563")
        ax.set_ylabel("Efficiency Index (%)", fontsize=10, color="#4b5563")
        ax.grid(True, linestyle=":", alpha=0.6)
        ax.legend()
        plt.tight_layout()
        plt.savefig(output_path, dpi=300)
        plt.close(fig)
        logger.info("Successfully created chart image at %s", output_path)
    except Exception as e:
        logger.error("Failed to create chart image: %s", e)


def create_table(doc: Document, headers: List[str], rows: List[List[str]], font_name: str = "Times New Roman", font_size: int = 11):
    if not headers:
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        cell.text = ""
        _set_cell_shading(cell, "D9F2EF")
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header.strip())
        set_font_style(run, font_name=font_name, font_size=font_size, bold=True, color=RGBColor(17, 94, 89))
    _set_repeat_table_header(table.rows[0])

    for row_data in rows:
        row_cells = table.add_row().cells
        for index in range(len(headers)):
            value = row_data[index] if index < len(row_data) else ""
            cell = row_cells[index]
            cell.text = ""
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_formatted_text(paragraph, value.strip(), font_name=font_name, font_size=font_size)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    logger.info("Successfully created table with %d rows", len(rows))


def render_table_block(doc: Document, table_lines: List[str], font_name: str = "Times New Roman"):
    parsed_rows = []
    for line in table_lines:
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        if parts:
            parsed_rows.append(parts)

    if not parsed_rows:
        return

    headers = parsed_rows[0]
    data_rows = parsed_rows[1:]
    if data_rows and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in data_rows[0]):
        data_rows = data_rows[1:]

    create_table(doc, headers, data_rows, font_name=font_name)


def _render_chart(doc: Document, stripped: str, section: str, font_name: str, font_size: int) -> None:
    chart_path = os.path.join(tempfile.gettempdir(), f"chart_{section}.png")
    create_chart_image(chart_path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(chart_path, width=Cm(14))

    caption_text = "Figure 1: Generated Chart Output"
    match = re.search(r"\[Chart:\s*(.*?)\]", stripped)
    if match:
        caption_text = f"Figure: {match.group(1)}"
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_caption = caption.add_run(caption_text)
    set_font_style(run_caption, font_name=font_name, font_size=max(font_size - 2, 9), italic=True)


def _render_markdown_block(
    doc: Document,
    section: str,
    text_block: str,
    font_name: str,
    font_size: int,
    alignment,
    first_line_indent_cm: float,
    line_spacing: float,
) -> None:
    lines = text_block.split("\n")
    
    current_block_type = None  # None, "paragraph", "list", "table", "block_math"
    accumulated_lines: List[str] = []
    
    def flush_current_block() -> None:
        nonlocal current_block_type, accumulated_lines
        if not accumulated_lines:
            return
        
        block_text = "\n".join(accumulated_lines)
        
        if current_block_type == "table":
            render_table_block(doc, accumulated_lines, font_name=font_name)
        elif current_block_type == "list":
            for line in accumulated_lines:
                list_match = _LIST_RE.match(line)
                if list_match:
                    marker = list_match.group(2)
                    body = list_match.group(3)
                    if marker[0].isdigit():
                        # Numbered list: render manually to avoid auto-incrementing across the document.
                        paragraph = doc.add_paragraph(style="Body Text")
                        paragraph.paragraph_format.left_indent = Cm(0.75)
                        paragraph.paragraph_format.first_line_indent = Cm(0.0)
                        paragraph.paragraph_format.line_spacing = line_spacing
                        paragraph.paragraph_format.space_after = Pt(3)

                        prefix_run = paragraph.add_run(f"{marker} ")
                        set_font_style(prefix_run, font_name=font_name, font_size=font_size)
                        add_formatted_text(paragraph, body, font_name=font_name, font_size=font_size)
                    else:
                        # Bullet list
                        paragraph = doc.add_paragraph(style="List Bullet")
                        paragraph.paragraph_format.space_after = Pt(3)
                        add_formatted_text(paragraph, body, font_name=font_name, font_size=font_size)
        elif current_block_type == "block_math":
            math_text = block_text.strip()
            if math_text.startswith("$$") and math_text.endswith("$$"):
                math_text = math_text[2:-2]
            elif math_text.startswith(r"\[") and math_text.endswith(r"\]"):
                math_text = math_text[2:-2]
            elif math_text.startswith(r"\["):
                math_text = math_text[2:]
            elif math_text.endswith(r"\]"):
                math_text = math_text[:-2]
                
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            parse_math_content(paragraph, math_text, font_name=font_name, font_size=font_size, italic=True)
        elif current_block_type == "paragraph":
            single_line_text = " ".join(l.strip() for l in accumulated_lines if l.strip())
            _add_markdown_paragraph(
                doc,
                single_line_text,
                "Body Text",
                font_name,
                font_size,
                alignment,
                first_line_indent_cm,
                line_spacing,
            )
            
        accumulated_lines = []
        current_block_type = None

    in_math_block = False
    math_delimiter = None
    
    for line in lines:
        stripped = line.strip()
        
        if not in_math_block:
            if stripped.startswith("$$"):
                flush_current_block()
                in_math_block = True
                math_delimiter = "$$"
                accumulated_lines.append(line)
                current_block_type = "block_math"
                if len(stripped) > 2 and stripped.endswith("$$"):
                    in_math_block = False
                    flush_current_block()
                continue
            elif stripped.startswith(r"\["):
                flush_current_block()
                in_math_block = True
                math_delimiter = r"\]"
                accumulated_lines.append(line)
                current_block_type = "block_math"
                if len(stripped) > 2 and stripped.endswith(r"\]"):
                    in_math_block = False
                    flush_current_block()
                continue
        else:
            accumulated_lines.append(line)
            if (math_delimiter == "$$" and stripped.endswith("$$")) or (math_delimiter == r"\]" and stripped.endswith(r"\]")):
                in_math_block = False
                flush_current_block()
            continue

        if not stripped:
            flush_current_block()
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            flush_current_block()
            level = min(len(heading.group(1)), 3)
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_formatted_text(paragraph, heading.group(2), font_name=font_name, font_size=font_size + 4 - level)
            continue

        if "[Chart]" in stripped or "[Chart:" in stripped:
            flush_current_block()
            _render_chart(doc, stripped, section, font_name, font_size)
            continue

        # Check for standard markdown image: ![alt text](image_path)
        img_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if img_match:
            flush_current_block()
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            if os.path.exists(img_path):
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(img_path, width=Cm(14))
                if alt_text:
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_caption = caption.add_run(f"Figure: {alt_text}")
                    set_font_style(run_caption, font_name=font_name, font_size=max(font_size - 2, 9), italic=True)
            else:
                logger.warning("Image path not found: %s", img_path)
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"[Image Not Found: {img_path}]")
                set_font_style(run, font_name=font_name, font_size=font_size, italic=True)
            continue

        if stripped.startswith("|"):
            if current_block_type != "table":
                flush_current_block()
                current_block_type = "table"
            accumulated_lines.append(line)
            continue

        list_match = _LIST_RE.match(line)
        if list_match:
            if current_block_type != "list":
                flush_current_block()
                current_block_type = "list"
            accumulated_lines.append(line)
            continue

        if current_block_type != "paragraph":
            flush_current_block()
            current_block_type = "paragraph"
        accumulated_lines.append(line)

    flush_current_block()


def render_paper(content: Dict[str, str], output_filename: str = "Output.docx", config: Optional[Any] = None):
    font_name = "Times New Roman"
    font_size = 14
    title_font_size = 20
    line_spacing = 1.5
    first_line_indent_cm = 1.25
    alignment_str = "justify"
    title_text = "GENERATED ACADEMIC PAPER"
    order = ["theory", "calculation", "conclusion"]

    if config is not None:
        if hasattr(config, "style") and config.style is not None:
            font_name = getattr(config.style, "font_name", font_name)
            font_size = getattr(config.style, "font_size", font_size)
            title_font_size = getattr(config.style, "title_font_size", title_font_size)
            line_spacing = getattr(config.style, "line_spacing", line_spacing)
            first_line_indent_cm = getattr(config.style, "first_line_indent_cm", first_line_indent_cm)
            alignment_str = getattr(config.style, "alignment", alignment_str)
        if hasattr(config, "pipeline") and config.pipeline is not None:
            title_text = getattr(config.pipeline, "title", title_text)
            if hasattr(config.pipeline, "sections") and config.pipeline.sections:
                order = [section.name for section in config.pipeline.sections]

    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    align_val = alignment_map.get(str(alignment_str).lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)

    output_dir = os.path.dirname(output_filename)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    _configure_styles(doc, font_name, font_size, title_font_size)

    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_paragraph(style="Document Title")
    title.paragraph_format.space_before = Pt(120)
    title_run = title.add_run(title_text)
    set_font_style(
        title_run,
        font_name=font_name,
        font_size=title_font_size,
        bold=True,
        color=RGBColor(15, 118, 110),
    )
    doc.add_page_break()

    for section_name in order:
        text_block = content.get(section_name, "")
        if not text_block.strip():
            continue
        _render_markdown_block(
            doc,
            section_name,
            text_block,
            font_name,
            font_size,
            align_val,
            first_line_indent_cm,
            line_spacing,
        )

    if not any(value.strip() for value in content.values()):
        paragraph = doc.add_paragraph(style="Body Text")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("No content was generated.")
        set_font_style(run, font_name=font_name, font_size=font_size, italic=True)

    try:
        doc.save(output_filename)
        logger.info("Saved document successfully to %s", output_filename)
    except Exception as exc:
        logger.error("Failed to save generated document to %s: %s", output_filename, exc)
        raise

    return output_filename
