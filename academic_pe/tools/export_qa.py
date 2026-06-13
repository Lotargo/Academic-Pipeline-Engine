from __future__ import annotations

import os
import ntpath
import re
import shutil
import subprocess
import tempfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document

from academic_pe.core.config import AppConfig
from academic_pe.tools.docx_renderer import render_paper
from academic_pe.tools.libreoffice import discover_soffice


@dataclass
class ExportIssue:
    severity: str
    message: str


@dataclass
class RenderResult:
    status: str
    soffice_path: Optional[str] = None
    pdf_path: Optional[str] = None
    png_paths: List[str] = field(default_factory=list)
    message: str = ""


@dataclass
class ExportResult:
    status: str
    filename: str
    path: str
    issues: List[ExportIssue]
    render: RenderResult

    def to_dict(self) -> dict:
        return {
            "status": self.status,
            "filename": self.filename,
            "path": self.path,
            "issues": [asdict(issue) for issue in self.issues],
            "render": asdict(self.render),
        }


_RAW_MARKDOWN_RE = re.compile(r"(^|\s)(#{1,6}\s|\*\*[^*]+\*\*|\$[^$]+\$|\$\$[^$]+\$\$)")
DEFAULT_EXPORT_FILENAME = "Final_Academic_Paper.docx"
_CONTROL_CHARS_RE = re.compile(r"[\x00-\x1f\x7f]")
_UNSAFE_FILENAME_CHARS_RE = re.compile(r'[*?"<>|]')
_WINDOWS_RESERVED_NAMES = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    *(f"COM{i}" for i in range(1, 10)),
    *(f"LPT{i}" for i in range(1, 10)),
}


def sanitize_filename(title: str) -> str:
    filename = _CONTROL_CHARS_RE.sub("", title or "")
    filename = filename.replace("\\", "-").replace("/", "-").replace(":", " -")
    filename = _UNSAFE_FILENAME_CHARS_RE.sub("", filename)
    filename = filename.strip(" .")
    if not filename:
        return "Untitled"

    stem = filename.rsplit(".", 1)[0]
    if stem.upper() in _WINDOWS_RESERVED_NAMES:
        filename = f"_{filename}"
    return filename


def resolve_export_filename(title: str, output_filename: Optional[str] = None, extension: str = ".docx") -> str:
    if not extension.startswith("."):
        extension = f".{extension}"
    filename = ntpath.basename(output_filename or DEFAULT_EXPORT_FILENAME)
    if filename == DEFAULT_EXPORT_FILENAME:
        filename = sanitize_filename(title)
    else:
        filename = sanitize_filename(filename)
    if not filename.lower().endswith(extension.lower()):
        stem, suffix = os.path.splitext(filename)
        filename = f"{stem or filename}{extension}" if suffix else f"{filename}{extension}"
    return filename


def convert_docx_to_pdf(docx_path: str, output_dir: str, output_filename: Optional[str] = None) -> RenderResult:
    discovery = discover_soffice()
    if not discovery.available or not discovery.executable:
        return RenderResult(
            status="skipped",
            message=f"LibreOffice/soffice not found. {discovery.install_hint}",
        )

    os.makedirs(output_dir, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp_dir:
        cmd = [
            discovery.executable,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            tmp_dir,
            docx_path,
        ]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120, errors="replace")
        if proc.returncode != 0:
            return RenderResult(
                status="failed",
                soffice_path=discovery.executable,
                message=(proc.stderr or proc.stdout or "LibreOffice conversion failed.").strip(),
            )

        pdf_candidates = list(Path(tmp_dir).glob("*.pdf"))
        if not pdf_candidates:
            return RenderResult(
                status="failed",
                soffice_path=discovery.executable,
                message="LibreOffice conversion did not produce a PDF.",
            )

        pdf_src = pdf_candidates[0]
        pdf_name = output_filename or pdf_src.name
        pdf_path = os.path.join(output_dir, os.path.basename(pdf_name))
        shutil.copyfile(pdf_src, pdf_path)

    return RenderResult(
        status="passed",
        soffice_path=discovery.executable,
        pdf_path=pdf_path,
        message="Converted DOCX to PDF with LibreOffice.",
    )


def inspect_docx_artifacts(docx_path: str, required_sections: List[str]) -> List[ExportIssue]:
    issues: List[ExportIssue] = []
    document = Document(docx_path)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    if not full_text.strip():
        issues.append(ExportIssue("error", "DOCX contains no paragraph text."))

    if _RAW_MARKDOWN_RE.search(full_text):
        issues.append(ExportIssue("error", "DOCX contains visible Markdown or LaTeX delimiter artifacts."))

    if "$" in full_text:
        issues.append(ExportIssue("error", "DOCX contains raw dollar signs from LaTeX formulas."))

    for section in required_sections:
        if section not in full_text and not any(section in getattr(paragraph.style, "name", "").lower() for paragraph in document.paragraphs):
            continue

    for table in document.tables:
        if not table.rows:
            issues.append(ExportIssue("warning", "DOCX contains an empty table."))

    return issues


def render_docx_pages(docx_path: str, qa_dir: str) -> RenderResult:
    convert_result = convert_docx_to_pdf(docx_path, qa_dir)
    if convert_result.status != "passed" or not convert_result.pdf_path:
        return convert_result

    pdf_path = convert_result.pdf_path

    try:
        import fitz  # type: ignore
    except ImportError:
        return RenderResult(
            status="partial",
            soffice_path=convert_result.soffice_path,
            pdf_path=pdf_path,
            message="PDF rendered, but PyMuPDF is not installed, so PNG page rendering was skipped.",
        )

    png_paths: List[str] = []
    pdf_doc = fitz.open(pdf_path)
    try:
        for page_index in range(pdf_doc.page_count):
            page = pdf_doc.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            png_path = os.path.join(qa_dir, f"page-{page_index + 1}.png")
            pix.save(png_path)
            png_paths.append(png_path)
    finally:
        pdf_doc.close()

    return RenderResult(
        status="passed" if png_paths else "failed",
        soffice_path=convert_result.soffice_path,
        pdf_path=pdf_path,
        png_paths=png_paths,
        message="Rendered DOCX to PNG pages." if png_paths else "PDF had no pages to render.",
    )


def export_docx_with_qa(content: Dict[str, str], config: AppConfig, output_filename: Optional[str] = None) -> ExportResult:
    output_dir = config.pipeline.output_dir
    os.makedirs(output_dir, exist_ok=True)

    requested_filename = output_filename or config.pipeline.output_filename
    filename = resolve_export_filename(config.pipeline.title, requested_filename)
    docx_path = os.path.join(output_dir, filename)
    render_paper(content, output_filename=docx_path, config=config)

    issues = inspect_docx_artifacts(
        docx_path,
        required_sections=[section.name for section in config.pipeline.sections],
    )

    qa_dir = os.path.join(output_dir, "_qa", Path(filename).stem)
    render_result = render_docx_pages(docx_path, qa_dir)
    if render_result.status == "failed":
        issues.append(ExportIssue("error", f"Render QA failed: {render_result.message}"))
    elif render_result.status in {"skipped", "partial"}:
        issues.append(ExportIssue("warning", render_result.message))

    status = "failed" if any(issue.severity == "error" for issue in issues) else "passed"
    return ExportResult(status=status, filename=filename, path=docx_path, issues=issues, render=render_result)


def export_pdf_with_qa(content: Dict[str, str], config: AppConfig, output_filename: Optional[str] = None) -> ExportResult:
    output_dir = config.pipeline.output_dir
    os.makedirs(output_dir, exist_ok=True)

    requested_filename = output_filename or config.pipeline.output_filename
    filename = resolve_export_filename(config.pipeline.title, requested_filename, extension=".pdf")
    pdf_path = os.path.join(output_dir, filename)

    issues: List[ExportIssue] = []
    with tempfile.TemporaryDirectory() as tmp_dir:
        docx_filename = resolve_export_filename(config.pipeline.title, requested_filename, extension=".docx")
        docx_path = os.path.join(tmp_dir, docx_filename)
        render_paper(content, output_filename=docx_path, config=config)

        issues.extend(
            inspect_docx_artifacts(
                docx_path,
                required_sections=[section.name for section in config.pipeline.sections],
            )
        )

        render_result = convert_docx_to_pdf(docx_path, output_dir, output_filename=filename)

    if render_result.status == "failed":
        issues.append(ExportIssue("error", f"PDF conversion failed: {render_result.message}"))
    elif render_result.status == "skipped":
        issues.append(ExportIssue("error", render_result.message))
    elif render_result.status == "passed":
        try:
            import fitz  # type: ignore
        except ImportError:
            issues.append(ExportIssue("warning", "PyMuPDF is not installed, so PDF structure validation was skipped."))
        else:
            if render_result.pdf_path:
                pdf_doc = fitz.open(render_result.pdf_path)
                try:
                    if pdf_doc.page_count == 0:
                        issues.append(ExportIssue("error", "PDF contains no pages."))
                finally:
                    pdf_doc.close()

    status = "failed" if any(issue.severity == "error" for issue in issues) else "passed"
    return ExportResult(status=status, filename=filename, path=pdf_path, issues=issues, render=render_result)
